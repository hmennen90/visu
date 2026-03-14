#!/usr/bin/env python3
"""Generate VISU Engine documentation as DOCX."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# -- Helper functions --
def add_code_block(text, doc=doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    # Background shading
    shading = run._element.get_or_add_rPr()
    s = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:fill'): 'F5F5F5',
    })
    shading.append(s)
    return p

def add_table(headers, rows, doc=doc):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacer
    return table

def add_bullet(text, bold_prefix=None, doc=doc):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('VISU')
run.font.size = Pt(48)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PHP Game Engine')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Technische Dokumentation')
run.font.size = Pt(16)

doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('Version 2.0 — März 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Inhaltsverzeichnis', level=1)
toc_items = [
    '1. Überblick',
    '2. Architektur',
    '3. Technischer Stack',
    '4. Projektstruktur',
    '5. Entity Component System (ECS)',
    '6. Scene System',
    '7. Render Pipeline',
    '8. FlyUI — Immediate-Mode GUI',
    '9. UI System (JSON-basiert)',
    '10. Audio System',
    '11. 2D Collision System',
    '12. Camera 2D System',
    '13. Save/Load System',
    '14. Signal & Event System',
    '15. Transpiler (Build-Optimierung)',
    '16. Input System',
    '17. Distribution',
    '18. Entwicklungsumgebung',
    '19. Roadmap',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. ÜBERBLICK
# ============================================================
doc.add_heading('1. Überblick', level=1)

doc.add_paragraph(
    'VISU ist eine PHP-native Game Engine für 2D- und 3D-Spiele. '
    'Sie basiert auf dem OpenGL-Framework php-glfw und erweitert dieses '
    'um ein vollständiges Entity Component System, Scene Management, '
    'Audio, UI, Collision Detection und mehr.'
)

doc.add_paragraph(
    'Die Engine ist als CLI-Anwendung konzipiert und nutzt einen Fixed-Timestep '
    'Game Loop mit variablem Rendering. Das Datenformat ist durchgehend JSON — '
    'für Szenen, UI-Layouts, Konfiguration und Spielstände.'
)

doc.add_heading('Projektziele', level=2)
add_bullet('2D Engine — Komplett und produktionsbereit', bold_prefix=None)
add_bullet('3D Engine — In Planung (Mesh-Loading, Lighting, Physics)', bold_prefix=None)
add_bullet('AI-gestütztes Game Authoring — Natürliche Sprache als primäres Authoring-Interface', bold_prefix=None)
add_bullet('Open Source — MIT-Lizenz', bold_prefix=None)

# ============================================================
# 2. ARCHITEKTUR
# ============================================================
doc.add_heading('2. Architektur', level=1)

doc.add_paragraph('VISU folgt einer geschichteten Architektur:')

add_code_block(
    'Authoring Layer         →  LLM (Claude Code) + Vue SPA Editor + Manuell\n'
    'Datenformat             →  JSON (Szenen, UI-Layouts, Config, Saves)\n'
    'Game-Agnostic Layer     →  Scene-System, UI-Interpreter, AudioManager,\n'
    '                           SaveManager, Mod-System\n'
    '────────────────────────────────────────────────────────────\n'
    'VISU Engine (src/)      →  ECS, Game Loop, Input, Render-Pipeline,\n'
    '  Namespace: VISU\\        Signal-System, Camera, FlyUI, Graphics,\n'
    '                           Shader Management, Font Rendering\n'
    '────────────────────────────────────────────────────────────\n'
    'C-Extensions            →  php-glfw (NanoVG 2D, OpenGL 4.1 3D)\n'
    'Hardware                →  GPU'
)

doc.add_heading('Kern-Entscheidungen', level=2)
add_table(
    ['Thema', 'Entscheidung'],
    [
        ['2D Rendering', 'php-glfw + NanoVG'],
        ['3D Rendering', 'php-glfw + OpenGL 4.1'],
        ['Editor UI', 'Vue SPA (Web-basiert, localhost)'],
        ['In-Game UI', 'JSON → UIInterpreter (kein HTML/CSS)'],
        ['Szenen-Format', 'JSON (Entity-Hierarchien, Transforms, Components)'],
        ['Save-Format', 'JSON (SaveManager mit Slots, Autosave, Migrationen)'],
        ['Audio-Format', 'WAV + MP3 (minimp3 via FFI)'],
        ['Game Loop', 'PHP CLI, Fixed Timestep + Variable Render'],
        ['Distribution', 'Launcher-Binary + statisches PHP-Binary + game.phar'],
        ['DI Container', 'ClanCats Container (.ctn Konfigurationsdateien)'],
    ]
)

# ============================================================
# 3. TECHNISCHER STACK
# ============================================================
doc.add_heading('3. Technischer Stack', level=1)

doc.add_heading('Voraussetzungen', level=2)
add_table(
    ['Komponente', 'Version', 'Beschreibung'],
    [
        ['PHP', '≥ 8.1 (CLI-SAPI)', 'JIT empfohlen für Performance'],
        ['php-glfw', '*', 'OpenGL 4.1 + NanoVG 2D (C-Extension)'],
        ['ext-ffi', 'optional', 'Für Audio (SDL3/OpenAL) und Gamepad-Support'],
        ['ClanCats Container', '^1.3', 'Dependency Injection'],
        ['League\\CLImate', '^3.8', 'CLI-Ausgabe'],
        ['Composer', 'aktuell', 'PSR-4 Autoloading (VISU\\ → src/)'],
    ]
)

doc.add_heading('Entwicklungswerkzeuge', level=2)
add_table(
    ['Tool', 'Zweck', 'Befehl'],
    [
        ['PHPUnit', 'Unit Tests (257+ Tests)', './vendor/bin/phpunit'],
        ['PHPStan', 'Statische Analyse (Level 8)', './vendor/bin/phpstan analyse'],
        ['PHPBench', 'Performance Benchmarks', './vendor/bin/phpbench run'],
    ]
)

doc.add_heading('Bootstrap & DI', level=2)
doc.add_paragraph('Die Engine nutzt Pfad-Konstanten, die vor dem Bootstrap definiert werden:')
add_table(
    ['Konstante', 'Beschreibung'],
    [
        ['VISU_PATH_ROOT', 'Projekt-Wurzelverzeichnis'],
        ['VISU_PATH_CACHE', 'Cache-Verzeichnis'],
        ['VISU_PATH_STORE', 'Persistenter Speicher'],
        ['VISU_PATH_RESOURCES', 'Anwendungs-Ressourcen'],
        ['VISU_PATH_APPCONFIG', 'Pfad zu anwendungsspezifischen .ctn Dateien'],
    ]
)

# ============================================================
# 4. PROJEKTSTRUKTUR
# ============================================================
doc.add_heading('4. Projektstruktur', level=1)

add_code_block(
    'visu/\n'
    '  composer.json              Package: phpgl/visu\n'
    '  visu.ctn                   DI-Container Konfiguration\n'
    '  bootstrap.php              Application Bootstrap\n'
    '  phpunit.xml                PHPUnit Konfiguration\n'
    '  phpstan.neon               PHPStan Level 8\n'
    '\n'
    '  src/                       Engine-Quellcode (Namespace: VISU\\)\n'
    '    Animation/               Transition-Animationen\n'
    '    Asset/                   AssetManager (Lazy-Loading + Caching)\n'
    '    Audio/                   AudioManager, Mp3Decoder, Backends\n'
    '    Command/                 CLI-Kommando-System\n'
    '    Component/               ECS-Components\n'
    '    ECS/                     Entity Component System\n'
    '    FlyUI/                   Immediate-Mode GUI\n'
    '    Geo/                     Geometrie (AABB, Ray, Frustum, Transform)\n'
    '    Graphics/                OpenGL-Rendering\n'
    '    Instrument/              Profiling\n'
    '    OS/                      Window, Input, GamepadManager\n'
    '    Runtime/                 GameLoop, DebugConsole\n'
    '    Save/                    SaveManager, SaveSlot\n'
    '    Scene/                   SceneLoader, SceneSaver, PrefabManager\n'
    '    SDL3/                    SDL3 FFI-Bindings\n'
    '    Signal/                  Event/Signal-System\n'
    '    System/                  ECS-Systeme (Camera2D, Collision2D, ...)\n'
    '    Transpiler/              JSON → PHP Build-Optimierung\n'
    '    UI/                      UIInterpreter, UIScreenStack\n'
    '\n'
    '  tests/                     PHPUnit Tests (spiegelt src/)\n'
    '  resources/                 Shader, Fonts, Models, Libraries\n'
    '  examples/                  Beispiel-Anwendungen\n'
    '  bin/visu                   CLI Entry Point'
)

# ============================================================
# 5. ECS
# ============================================================
doc.add_heading('5. Entity Component System (ECS)', level=1)

doc.add_paragraph(
    'Das ECS ist das Herzstück der Engine. Entities sind leichtgewichtige IDs, '
    'Components sind reine Datenobjekte, und Systems enthalten die Logik.'
)

doc.add_heading('Kernklassen', level=2)
add_table(
    ['Klasse', 'Pfad', 'Beschreibung'],
    [
        ['EntityRegistry', 'src/ECS/EntityRegistry.php', 'Verwaltet Entities mit Freelist-Pooling für effiziente Wiederverwendung'],
        ['ComponentRegistry', 'src/ECS/ComponentRegistry.php', 'Registriert und löst Component-Typen auf (String → Klasse)'],
        ['SystemInterface', 'src/ECS/SystemInterface.php', 'Interface für alle ECS-Systeme (register/unregister)'],
    ]
)

doc.add_heading('Verfügbare Components', level=2)
add_table(
    ['Component', 'Beschreibung'],
    [
        ['SpriteRenderer', 'Sprite-Darstellung mit Texture, Layer, Flip'],
        ['SpriteAnimator', 'Frame-basierte Sprite-Animation'],
        ['Tilemap', 'Tile-basierte Karten mit Auto-Tiling (Bitmask)'],
        ['NameComponent', 'Benennung von Entities'],
        ['BoxCollider2D', 'Rechteckiger 2D-Collider (AABB)'],
        ['CircleCollider2D', 'Kreisförmiger 2D-Collider'],
        ['AnimationComponent', 'Animations-Steuerung'],
        ['DirectionalLightComponent', 'Richtungslicht für 3D-Szenen'],
    ]
)

doc.add_heading('Beispiel: Entity erstellen', level=2)
add_code_block(
    '$entity = $entities->create();\n'
    '$sprite = $entities->attach($entity, new SpriteRenderer());\n'
    '$sprite->sprite = "assets/sprites/player.png";\n'
    '$sprite->sortingLayer = 10;\n'
    '\n'
    '$name = $entities->attach($entity, new NameComponent("Player"));'
)

# ============================================================
# 6. SCENE SYSTEM
# ============================================================
doc.add_heading('6. Scene System', level=1)

doc.add_paragraph(
    'Szenen werden als JSON-Dateien definiert und enthalten Entity-Hierarchien '
    'mit Transforms, Components und Kinder-Entities. Der SceneLoader/SceneSaver '
    'konvertiert bidirektional zwischen JSON und ECS-Entities.'
)

doc.add_heading('Kernklassen', level=2)
add_table(
    ['Klasse', 'Pfad', 'Beschreibung'],
    [
        ['SceneLoader', 'src/Scene/SceneLoader.php', 'Lädt JSON-Szenen in EntityRegistry'],
        ['SceneSaver', 'src/Scene/SceneSaver.php', 'Serialisiert Entities zurück nach JSON'],
        ['SceneManager', 'src/Scene/SceneManager.php', 'Verwaltet aktive Szenen und Übergänge'],
        ['PrefabManager', 'src/Scene/PrefabManager.php', 'Prefab-Instanziierung aus JSON-Templates'],
    ]
)

doc.add_heading('Szenen-Format', level=2)
add_code_block(
    '{\n'
    '  "entities": [{\n'
    '    "name": "Player",\n'
    '    "transform": {\n'
    '      "position": [0, 1, 0],\n'
    '      "rotation": [0, 0, 0],\n'
    '      "scale": [1, 1, 1]\n'
    '    },\n'
    '    "components": [\n'
    '      { "type": "SpriteRenderer", "sprite": "assets/player.png" },\n'
    '      { "type": "CharacterController", "speed": 5.0 }\n'
    '    ],\n'
    '    "children": []\n'
    '  }]\n'
    '}'
)

# ============================================================
# 7. RENDER PIPELINE
# ============================================================
doc.add_heading('7. Render Pipeline', level=1)

doc.add_paragraph(
    'Die Render Pipeline basiert auf OpenGL 4.1 via php-glfw. '
    'Für 2D-Rendering wird NanoVG verwendet, für 3D steht die volle '
    'OpenGL-Pipeline mit Shader-Management, Framebuffers und Deferred Rendering zur Verfügung.'
)

doc.add_heading('Kernmodule', level=2)
add_table(
    ['Modul', 'Pfad', 'Beschreibung'],
    [
        ['ShaderProgram', 'src/Graphics/', 'Shader-Kompilierung und -Verwaltung'],
        ['Framebuffer', 'src/Graphics/', 'Off-Screen Rendering, GBuffer'],
        ['Texture', 'src/Graphics/', 'Texture-Loading und -Management'],
        ['RenderPipeline', 'src/Graphics/Rendering/', 'Multi-Pass Rendering'],
        ['SpriteBatchPass', 'src/Graphics/Rendering/', 'Batched 2D-Sprite-Rendering'],
        ['Camera', 'src/Graphics/', 'Kamera-Verwaltung (2D/3D)'],
        ['Font', 'src/Graphics/Font/', 'Bitmap Font Rendering'],
    ]
)

# ============================================================
# 8. FlyUI
# ============================================================
doc.add_heading('8. FlyUI — Immediate-Mode GUI', level=1)

doc.add_paragraph(
    'FlyUI ist VISUs eingebautes Immediate-Mode GUI-System. '
    'Widgets werden pro Frame aufgerufen — es gibt keinen persistenten Widget-Tree. '
    'Das System unterstützt Layout-Management, Theming und Input-Handling.'
)

doc.add_heading('Verfügbare Widgets', level=2)
add_table(
    ['Widget', 'Klasse', 'Beschreibung'],
    [
        ['Button', 'FUIButton', 'Klickbare Schaltfläche mit Label'],
        ['ButtonGroup', 'FUIButtonGroup', 'Gruppierte Schaltflächen'],
        ['Label/Text', 'FUIText / FUILabel', 'Textanzeige'],
        ['Card', 'FUICard', 'Container mit optionalem Titel'],
        ['Checkbox', 'FUICheckbox', 'Toggle-Schalter'],
        ['Select', 'FUISelect', 'Dropdown-Auswahl'],
        ['ProgressBar', 'FUIProgressBar', 'Fortschrittsbalken'],
        ['Space', 'FUISpace', 'Abstandshalter'],
        ['Layout', 'FUILayout', 'Container mit Flow-Richtung'],
    ]
)

doc.add_heading('Layout-System', level=2)
add_table(
    ['Enum', 'Werte', 'Beschreibung'],
    [
        ['FUILayoutFlow', 'Row, Column', 'Anordnungsrichtung'],
        ['FUILayoutAlignment', 'Start, Center, End, Stretch', 'Ausrichtung'],
        ['FUILayoutSizing', 'Fixed, Grow, Shrink', 'Größenverhalten'],
    ]
)

# ============================================================
# 9. UI SYSTEM
# ============================================================
doc.add_heading('9. UI System (JSON-basiert)', level=1)

doc.add_paragraph(
    'Für spielspezifische UIs bietet VISU ein JSON-basiertes UI-System. '
    'JSON-Layouts werden vom UIInterpreter rekursiv in FlyUI-Aufrufe übersetzt. '
    'Data Binding und Event Handling sind eingebaut.'
)

doc.add_heading('UI-Format', level=2)
add_code_block(
    '{\n'
    '  "type": "panel",\n'
    '  "layout": "column",\n'
    '  "padding": 10,\n'
    '  "children": [\n'
    '    { "type": "label", "text": "Geld: {economy.money}", "fontSize": 16 },\n'
    '    { "type": "progressbar", "value": "{player.health}", "color": "#0088ff" },\n'
    '    { "type": "button", "label": "Aktion", "event": "ui.action" }\n'
    '  ]\n'
    '}'
)

doc.add_heading('Node-Typen', level=2)
add_table(
    ['Typ', 'Beschreibung'],
    [
        ['panel', 'Container mit Layout (row/column)'],
        ['label', 'Textanzeige mit Data Binding ({path.to.value})'],
        ['button', 'Klickbar, löst UIEventSignal aus'],
        ['progressbar', 'Wertanzeige 0.0–1.0 mit Farbe'],
        ['checkbox', 'Toggle mit Event'],
        ['select', 'Dropdown-Auswahl'],
        ['image', 'Bildanzeige'],
        ['space', 'Abstandshalter'],
    ]
)

doc.add_heading('Kernklassen', level=2)
add_table(
    ['Klasse', 'Beschreibung'],
    [
        ['UIInterpreter', 'Wandelt JSON-Bäume in FlyUI-Aufrufe um'],
        ['UIDataContext', 'Stellt Daten für {path} Bindings bereit'],
        ['UIScreenStack', 'Stack-basierte Screen-Verwaltung (push/pop/replace)'],
        ['UIScreen', 'Einzelner UI-Screen mit optionaler Transparenz'],
        ['UITransition', 'Animierte Übergänge (FadeIn, SlideIn, ScaleIn, etc.)'],
        ['UIEventSignal', 'Signal für Button/Checkbox/Select Events'],
    ]
)

# ============================================================
# 10. AUDIO SYSTEM
# ============================================================
doc.add_heading('10. Audio System', level=1)

doc.add_paragraph(
    'Das Audio-System unterstützt WAV- und MP3-Dateien mit automatischer '
    'Backend-Erkennung. Zwei Backends stehen zur Verfügung: SDL3 (primär) '
    'und OpenAL (Fallback). MP3-Decoding erfolgt über minimp3 via PHP FFI.'
)

doc.add_heading('Architektur', level=2)
add_code_block(
    'AudioManager\n'
    '  ├── AudioBackendInterface\n'
    '  │     ├── SDL3AudioBackend    (SDL3 via FFI)\n'
    '  │     └── OpenALAudioBackend  (OpenAL via FFI)\n'
    '  ├── Mp3Decoder               (minimp3 via FFI)\n'
    '  └── AudioClipData            (Backend-agnostische PCM-Daten)'
)

doc.add_heading('Unterstützte Formate', level=2)
add_table(
    ['Format', 'Decoder', 'Beschreibung'],
    [
        ['WAV', 'Backend-nativ', 'Unkomprimiertes PCM, von beiden Backends direkt geladen'],
        ['MP3', 'minimp3 (FFI)', 'Komprimiert, via minimp3 C-Library dekodiert'],
    ]
)

doc.add_heading('MP3-Support (minimp3)', level=2)
doc.add_paragraph(
    'minimp3 ist eine Header-only C-Library, die als Shared Library '
    'vorkompiliert für alle Plattformen mitgeliefert wird:'
)
add_table(
    ['Plattform', 'Pfad', 'Größe'],
    [
        ['macOS arm64', 'resources/lib/minimp3/darwin-arm64/libminimp3.dylib', '~84 KB'],
        ['macOS x86_64', 'resources/lib/minimp3/darwin-x86_64/libminimp3.dylib', '~48 KB'],
        ['Linux x86_64', 'resources/lib/minimp3/linux-x86_64/libminimp3.so', '~120 KB'],
        ['Windows x86_64', 'resources/lib/minimp3/windows-x86_64/minimp3.dll', '~196 KB'],
    ]
)

doc.add_heading('API-Beispiele', level=2)
add_code_block(
    '// AudioManager erstellen (Auto-Detection)\n'
    '$audio = AudioManager::create($sdl);\n'
    '\n'
    '// Sound-Effekte abspielen (WAV oder MP3)\n'
    '$audio->playSound("assets/sfx/explosion.wav");\n'
    '$audio->playSound("assets/sfx/coin.mp3");\n'
    '\n'
    '// Musik abspielen (automatisches Looping)\n'
    '$audio->playMusic("assets/music/theme.mp3");\n'
    '$audio->stopMusic();\n'
    '\n'
    '// Kanal-Lautstärke\n'
    '$audio->setChannelVolume(AudioChannel::Music, 0.7);\n'
    '$audio->setChannelVolume(AudioChannel::SFX, 1.0);\n'
    '\n'
    '// Im Game Loop aufrufen (für Music-Looping)\n'
    '$audio->update();'
)

# ============================================================
# 11. COLLISION SYSTEM
# ============================================================
doc.add_heading('11. 2D Collision System', level=1)

doc.add_paragraph(
    'Das 2D-Kollisionssystem verwendet ein Spatial Grid für die Broad Phase '
    'und AABB/Circle-Tests für die Narrow Phase. Es unterstützt Trigger-Events '
    '(ENTER/STAY/EXIT) und Layer-basierte Filterung.'
)

doc.add_heading('Components', level=2)
add_table(
    ['Component', 'Beschreibung'],
    [
        ['BoxCollider2D', 'Rechteckiger Collider (AABB), konfigurierbare Größe und Offset'],
        ['CircleCollider2D', 'Kreisförmiger Collider mit Radius und Offset'],
    ]
)

doc.add_heading('Features', level=2)
add_bullet('Spatial Grid Broad Phase — Effiziente Vorauswahl potenzieller Kollisionspaare')
add_bullet('AABB/Circle Narrow Phase — Präzise Kollisionserkennung')
add_bullet('Trigger ENTER/STAY/EXIT — Zustandsbasierte Events')
add_bullet('CollisionSignal / TriggerSignal — Integration ins Signal-System')
add_bullet('Layer/Mask Bitmask — Feingranulare Kollisionsfilterung')
add_bullet('Raycast2D — Point Query und Ray Cast mit maxDistance und Layer-Filter')

# ============================================================
# 12. CAMERA 2D
# ============================================================
doc.add_heading('12. Camera 2D System', level=1)

doc.add_paragraph(
    'Das Camera2DSystem bietet eine vollständige 2D-Kamera mit Follow-Mechanik, '
    'Begrenzungen, Zoom und Screen-Shake.'
)

doc.add_heading('Features', level=2)
add_table(
    ['Feature', 'Beschreibung'],
    [
        ['Follow Target', 'Kamera folgt einem Entity mit konfigurierbarem Offset'],
        ['Smooth Damping', 'Weiche Kamerabewegung mit Dämpfungsfaktor'],
        ['Bounds', 'Kamera bleibt innerhalb definierter Grenzen'],
        ['Zoom', 'Stufenloses Zoomen'],
        ['Shake', 'Bildschirmwackeln mit konfigurierbarer Intensität und Dauer'],
    ]
)

# ============================================================
# 13. SAVE/LOAD
# ============================================================
doc.add_heading('13. Save/Load System', level=1)

doc.add_paragraph(
    'Der SaveManager verwaltet Spielstände als JSON-Dateien in benannten Slots. '
    'Er unterstützt Autosave, Schema-Versionierung und Daten-Migration.'
)

doc.add_heading('Kernklassen', level=2)
add_table(
    ['Klasse', 'Beschreibung'],
    [
        ['SaveManager', 'Hauptklasse: Save/Load/Delete, Autosave, Migrationen'],
        ['SaveSlot', 'Einzelner Spielstand: gameState + sceneData + Metadaten'],
        ['SaveSlotInfo', 'Kompakte Metadaten (Timestamp, PlayTime, Description)'],
    ]
)

doc.add_heading('Features', level=2)
add_bullet('Benannte Slots — Mehrere Spielstände parallel')
add_bullet('Autosave — Konfigurierbares Intervall und Slot-Name')
add_bullet('Schema-Versionierung — Versions-Nummer pro Spielstand')
add_bullet('Migration-System — registerMigration() für Daten-Upgrades')
add_bullet('Slot-Listing — Nach Timestamp sortiert')
add_bullet('Sicherheit — Slot-Name Sanitization gegen Directory Traversal')
add_bullet('Events — SaveSignal (save.completed, save.loaded, save.deleted)')

# ============================================================
# 14. SIGNAL SYSTEM
# ============================================================
doc.add_heading('14. Signal & Event System', level=1)

doc.add_paragraph(
    'VISU nutzt ein typsicheres Signal-System für lose Kopplung zwischen '
    'Engine-Modulen und Spiellogik. Der Dispatcher ermöglicht das Registrieren '
    'und Auslösen von Events.'
)

doc.add_heading('Vordefinierte Signale', level=2)
add_table(
    ['Kategorie', 'Signale', 'Beschreibung'],
    [
        ['Bootstrap', 'BootstrapSignal', 'Engine-Start und -Initialisierung'],
        ['Input', 'InputSignal', 'Tastatur, Maus, Gamepad Events'],
        ['ECS', 'EntitySpawnedSignal, EntityDestroyedSignal', 'Entity-Lifecycle'],
        ['Scene', 'SceneLoadedSignal, SceneUnloadedSignal', 'Szenen-Wechsel'],
        ['Collision', 'CollisionSignal, TriggerSignal', 'Kollisionen und Trigger'],
        ['UI', 'UIEventSignal', 'Button/Checkbox/Select Interaktionen'],
        ['Save', 'SaveSignal', 'Spielstand gespeichert/geladen/gelöscht'],
    ]
)

doc.add_heading('Beispiel', level=2)
add_code_block(
    '// Signal registrieren\n'
    '$dispatcher->register(CollisionSignal::class, function(CollisionSignal $sig) {\n'
    '    echo "Kollision: {$sig->entityA} ↔ {$sig->entityB}";\n'
    '});\n'
    '\n'
    '// Signal auslösen\n'
    '$dispatcher->dispatch(new CollisionSignal($entityA, $entityB));'
)

# ============================================================
# 15. TRANSPILER
# ============================================================
doc.add_heading('15. Transpiler (Build-Optimierung)', level=1)

doc.add_paragraph(
    'Der Transpiler wandelt JSON-Definitionen in optimierte PHP-Factory-Klassen um. '
    'JSON bleibt Source of Truth (für den Editor), die generierten PHP-Klassen '
    'dienen als Laufzeit-Cache für maximale Performance.'
)

add_code_block(
    'JSON (Editor, editierbar)  ──transpile──→  PHP-Factory (Laufzeit, schnell)')

doc.add_heading('Transpiler-Typen', level=2)
add_table(
    ['Transpiler', 'Input', 'Output', 'Beschreibung'],
    [
        ['SceneTranspiler', 'scenes/*.json', 'Generated\\Scenes\\*.php', 'Direkte Entity-Erzeugung ohne Reflection'],
        ['UITranspiler', 'ui/*.json', 'Generated\\UI\\*.php', 'Direkte FlyUI-Aufrufe statt JSON-Parsing'],
        ['PrefabTranspiler', 'prefabs/*.json', 'Generated\\Prefabs\\*.php', 'Delegiert an SceneTranspiler'],
        ['TranspilerRegistry', '—', '—', 'MD5-Hashing für inkrementelle Builds'],
    ]
)

doc.add_heading('Performance-Vorteile', level=2)
add_bullet('Kein json_decode() zur Laufzeit')
add_bullet('Keine ComponentRegistry String→Klasse Auflösung')
add_bullet('Keine Reflection-basierte Property-Zuweisung')
add_bullet('Keine dynamische Typ-Konvertierung')
add_bullet('PHP Opcache-optimierbar und statisch analysierbar')

# ============================================================
# 16. INPUT
# ============================================================
doc.add_heading('16. Input System', level=1)

doc.add_paragraph(
    'Das Input-System abstrahiert Tastatur, Maus und Gamepad-Eingaben. '
    'InputActionMaps ermöglichen die Zuordnung von Aktionen zu beliebigen '
    'Eingabequellen.'
)

add_table(
    ['Modul', 'Beschreibung'],
    [
        ['InputActionMap', 'Abstrakte Aktionen (z.B. "jump") → physische Tasten'],
        ['InputContextMap', 'Kontext-abhängige Input-Profile (Menu, Gameplay, etc.)'],
        ['GamepadManager', 'SDL3-basierte Gamepad-Unterstützung via FFI'],
        ['Key / MouseButton', 'Enums für Tastatur- und Maustasten'],
    ]
)

# ============================================================
# 17. DISTRIBUTION
# ============================================================
doc.add_heading('17. Distribution', level=1)

doc.add_paragraph(
    'VISU-Spiele werden als selbstständige Pakete verteilt. '
    'Ein Launcher-Binary startet ein statisch kompiliertes PHP-Binary '
    'mit einem game.phar Archiv.'
)

add_code_block(
    'game_name/\n'
    '  game_name           ← Launcher-Binary\n'
    '  runtime/php         ← Statisches PHP 8.3 Binary (~15–25 MB)\n'
    '  game.phar           ← Engine + Game Logic (Opcache-geschützt)\n'
    '  assets/             ← Sprites, Sounds, UI-JSONs, Szenen\n'
    '  saves/              ← Spielstände (SaveManager)\n'
    '  mods/               ← Offen für Modder'
)

# ============================================================
# 18. ENTWICKLUNG
# ============================================================
doc.add_heading('18. Entwicklungsumgebung', level=1)

doc.add_heading('Setup', level=2)
add_code_block(
    'git clone <repository>\n'
    'cd visu\n'
    'composer install'
)

doc.add_heading('Tests ausführen', level=2)
add_code_block(
    '# Alle Tests\n'
    './vendor/bin/phpunit\n'
    '\n'
    '# Einzelner Test\n'
    './vendor/bin/phpunit --filter Mp3DecoderTest\n'
    '\n'
    '# Statische Analyse\n'
    './vendor/bin/phpstan analyse'
)

doc.add_heading('minimp3 neu kompilieren (optional)', level=2)
add_code_block(
    '# Mit zig (alle Plattformen):\n'
    'cd resources/lib/minimp3\n'
    './build.sh\n'
    '\n'
    '# Oder nativ (nur aktuelle Plattform):\n'
    'cc -shared -O2 -fPIC -o libminimp3.dylib minimp3_wrapper.c'
)

# ============================================================
# 19. ROADMAP
# ============================================================
doc.add_heading('19. Roadmap', level=1)

add_table(
    ['Phase', 'Status', 'Beschreibung'],
    [
        ['Phase 1 — Engine-Kern & Scenes', 'Abgeschlossen', 'ECS, SceneLoader, Sprites, Tilemap, Camera, Signals'],
        ['Phase 2 — Interaktion & UI', 'Abgeschlossen', '2D Collision, UI System, Audio (WAV+MP3), Camera Shake, Save/Load'],
        ['Phase 3 — Transpiler', 'In Arbeit', 'JSON→PHP Transpiler für Scenes, UI, Prefabs; CLI-Tooling ausstehend'],
        ['Phase 4 — 3D Grundlagen', 'Offen', 'Mesh-Loading, 3D Camera, Lighting, Materials, 3D Collision, Physics'],
        ['Phase 5 — Erweiterte 3D', 'Offen', 'Partikel, Skeletal Animation, Terrain, Post-Processing, AI, Pathfinding'],
        ['Editor — Vue SPA', 'Offen', 'Editor-Server, Scene Hierarchy, Property Inspector, UI Layout Editor'],
    ]
)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'VISU_Engine_Documentation.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
